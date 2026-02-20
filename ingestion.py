import os
import io
from typing import List, Dict, Any

from utils import ensure_dirs, char_chunk_text, build_chunk_metadata

# File parsing dependencies (use lightweight fallbacks)
try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx  # python-docx
except Exception:
    docx = None


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def parse_txt(path: str) -> List[Dict[str, Any]]:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    # single page None for TXT
    return [{"text": text, "metadata": {"source": os.path.basename(path)}}]


def parse_pdf(path: str) -> List[Dict[str, Any]]:
    pages: List[Dict[str, Any]] = []
    if PyPDF2 is None:
        raise RuntimeError("PyPDF2 not installed. Install it to parse PDFs.")
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for i, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            pages.append({
                "text": text,
                "metadata": {"source": os.path.basename(path), "page": i}
            })
    return pages


def parse_docx(path: str) -> List[Dict[str, Any]]:
    if docx is None:
        raise RuntimeError("python-docx not installed. Install it to parse DOCX.")
    document = docx.Document(path)
    text_buf = io.StringIO()
    for p in document.paragraphs:
        text_buf.write(p.text + "\n")
    text = text_buf.getvalue()
    # No page; include title if available (first paragraph heuristic)
    title = document.paragraphs[0].text.strip() if document.paragraphs else None
    meta = {"source": os.path.basename(path)}
    if title:
        meta["title"] = title
    return [{"text": text, "metadata": meta}]


def parse_file(path: str) -> List[Dict[str, Any]]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(path)
    if ext == ".docx":
        return parse_docx(path)
    if ext == ".txt":
        return parse_txt(path)
    raise ValueError(f"Unsupported file type: {ext}")


def chunk_records(records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    chunks: List[Dict[str, Any]] = []
    cid = 1
    for rec in records:
        text = rec.get("text", "")
        base_meta = rec.get("metadata", {})
        for ch in char_chunk_text(text):
            meta = base_meta.copy()
            meta.update(build_chunk_metadata(meta.get("source", "unknown"), meta.get("page"), cid))
            chunks.append({"text": ch, "metadata": meta})
            cid += 1
    return chunks


def ingest(paths: List[str]) -> List[Dict[str, Any]]:
    ensure_dirs()
    all_chunks: List[Dict[str, Any]] = []
    for p in paths:
        records = parse_file(p)
        all_chunks.extend(chunk_records(records))
    return all_chunks
